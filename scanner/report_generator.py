"""
Professional DOCX Penetration Test Report Generator
Generates a polished pentest report from DAST scan results.
"""
import io
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SEVERITY_COLORS = {
    'critical': RGBColor(0x9B, 0x1C, 0x31),
    'high':     RGBColor(0xDC, 0x26, 0x26),
    'medium':   RGBColor(0xEA, 0x58, 0x0C),
    'low':      RGBColor(0xCA, 0x8A, 0x04),
    'info':     RGBColor(0x25, 0x63, 0xEB),
}

SEVERITY_BG = {
    'critical': 'F5C6CB',
    'high':     'FDCFCF',
    'medium':   'FFE0B2',
    'low':      'FFF9C4',
    'info':     'BBDEFB',
}

SEVERITY_ORDER = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3, 'info': 4}
SEVERITY_EVIDENCE_PREFIX = {'critical': 'C', 'high': 'H', 'medium': 'M', 'low': 'L', 'info': 'I'}

TABLE_HEADER_BG = '1F2937'
TABLE_ALT_ROW = 'F3F4F6'
CODE_BLOCK_BG = 'F1F5F9'

REMEDIATION = {
    'xss': (
        "Implement context-aware output encoding for all user-supplied data. "
        "Use Content-Security-Policy headers to mitigate impact. "
        "Apply input validation using allowlists where possible. "
        "Use frameworks with built-in XSS protection (e.g., React, Angular)."
    ),
    'sql_injection': (
        "Use parameterized queries or prepared statements for all database interactions. "
        "Apply the principle of least privilege to database accounts. "
        "Implement input validation and use an ORM where appropriate. "
        "Deploy a Web Application Firewall (WAF) as an additional layer."
    ),
    'command_injection': (
        "Avoid passing user input directly to system commands. "
        "Use language-specific APIs instead of shell commands. "
        "If shell execution is unavoidable, use strict allowlists and escape all input. "
        "Run applications with minimal OS privileges."
    ),
    'ssrf': (
        "Validate and sanitize all user-supplied URLs. "
        "Implement allowlists for permitted domains and IP ranges. "
        "Block requests to internal/private IP ranges (10.x, 172.16.x, 192.168.x, 127.x). "
        "Use a dedicated proxy for outbound requests with strict filtering."
    ),
    'open_redirect': (
        "Validate redirect URLs against an allowlist of permitted domains. "
        "Avoid using user-supplied input directly in redirect targets. "
        "Use relative URLs for internal redirects. "
        "Display a warning page before redirecting to external URLs."
    ),
    'path_traversal': (
        "Validate file paths against an allowlist of permitted directories. "
        "Use canonical path checking to prevent directory traversal sequences. "
        "Run the application in a chroot or container with limited filesystem access. "
        "Never pass raw user input to filesystem APIs."
    ),
    'lfi': (
        "Avoid including files based on user-controlled input. "
        "If dynamic inclusion is required, use a strict allowlist of permitted files. "
        "Disable remote file inclusion in PHP (allow_url_include=Off). "
        "Implement proper access controls on the filesystem."
    ),
    'idor': (
        "Implement proper authorization checks for every resource access. "
        "Use indirect references (e.g., mapping user-specific IDs) instead of direct database IDs. "
        "Apply the principle of least privilege to API endpoints. "
        "Log and monitor access patterns for anomalies."
    ),
    'csrf': (
        "Implement anti-CSRF tokens for all state-changing operations. "
        "Use the SameSite cookie attribute (Strict or Lax). "
        "Verify the Origin and Referer headers on the server side. "
        "Require re-authentication for sensitive actions."
    ),
    'header_injection': (
        "Sanitize all user input used in HTTP response headers. "
        "Strip or encode carriage return and line feed characters (CR/LF). "
        "Use framework-provided methods for setting response headers. "
        "Implement Content-Security-Policy and X-Content-Type-Options headers."
    ),
    'information_disclosure': (
        "Remove verbose error messages, stack traces, and debug information in production. "
        "Configure custom error pages that do not reveal internal details. "
        "Remove unnecessary HTTP response headers (Server, X-Powered-By). "
        "Conduct regular reviews of exposed information."
    ),
    'clickjacking': (
        "Set the X-Frame-Options header to DENY or SAMEORIGIN. "
        "Implement Content-Security-Policy frame-ancestors directive. "
        "Use JavaScript frame-busting as an additional defense layer."
    ),
    'cors': (
        "Configure Access-Control-Allow-Origin to specific trusted domains, never use wildcard (*) with credentials. "
        "Validate the Origin header server-side before reflecting it. "
        "Restrict Access-Control-Allow-Methods and Access-Control-Allow-Headers to necessary values."
    ),
    'jwt': (
        "Use strong signing algorithms (RS256 or ES256), never 'none'. "
        "Validate all JWT claims (exp, iss, aud) server-side. "
        "Store tokens securely and implement proper token revocation. "
        "Keep token lifetimes short and use refresh tokens."
    ),
    'rfi': (
        "Disable remote file inclusion (allow_url_include=Off in PHP). "
        "Validate all file paths against strict allowlists. "
        "Implement input validation that blocks URLs and remote paths. "
        "Use a WAF to detect and block remote inclusion attempts."
    ),
    'security_headers': (
        "Implement all recommended security headers: Content-Security-Policy, "
        "X-Content-Type-Options, X-Frame-Options, Strict-Transport-Security, "
        "Referrer-Policy, and Permissions-Policy. "
        "Regularly audit response headers using automated tools."
    ),
}

DEFAULT_REMEDIATION = (
    "Conduct a detailed manual review of this finding. "
    "Apply the principle of least privilege and defense in depth. "
    "Implement input validation, output encoding, and proper access controls. "
    "Consult OWASP guidelines for specific remediation steps."
)

_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DEFAULT_TEMPLATE_PATH = os.path.join(_PROJECT_ROOT, 'report_templates', 'default_pentest_template.docx')


def _finding_family(vuln_type):
    """Normalize noisy agent labels into report-level finding families."""
    vt = (vuln_type or 'finding').lower().replace('-', '_').replace(' ', '_')
    if vt in {'cookie_security_flags', 'session_cookie_misconfiguration', 'jwt_cookie_flags'}:
        return 'insecure_jwt_cookie_flags'
    if 'alg_none' in vt:
        return 'jwt_alg_none_bypass'
    if 'jwk' in vt:
        return 'jwt_embedded_jwk_trust'
    if 'signature' in vt and 'jwt' in vt:
        return 'jwt_signature_bypass'
    return vt


def _display_finding_name(vuln_type):
    vt = _finding_family(vuln_type)
    names = {
        'insecure_jwt_cookie_flags': 'Insecure JWT Cookie Attributes',
        'jwt_alg_none_bypass': 'JWT Algorithm None Bypass',
        'jwt_embedded_jwk_trust': 'JWT Embedded JWK Trust',
        'jwt_signature_bypass': 'JWT Signature Validation Bypass',
    }
    return names.get(vt, vt.replace('_', ' ').title())


def _unique_ordered(values):
    out = []
    seen = set()
    for value in values:
        value = str(value or '').strip()
        if value and value not in seen:
            seen.add(value)
            out.append(value)
    return out


def _consolidate_findings(findings):
    """Collapse duplicate agent events into report-ready findings.

    The scan database intentionally stores every observed agent finding. A DOCX
    report should not. This groups the same vulnerability family/severity into
    one finding and lists affected endpoints as evidence.
    """
    grouped = {}
    for f in findings or []:
        sev = (f.get('severity') or 'info').lower()
        family = _finding_family(f.get('vuln_type'))
        key = (sev, family)
        entry = grouped.get(key)
        if not entry:
            entry = dict(f)
            entry['vuln_type'] = family
            entry['display_name'] = _display_finding_name(f.get('vuln_type'))
            entry['_evidence_items'] = []
            entry['_affected_urls'] = []
            entry['_parameters'] = []
            grouped[key] = entry

        entry['_affected_urls'].append(f.get('url'))
        entry['_parameters'].append(f.get('parameter'))
        evidence = (f.get('evidence') or '').strip()
        if evidence:
            entry['_evidence_items'].append(evidence)

        # Prefer the first finding that has a real replay command/response.
        for field in ('curl_command', 'request', 'response', 'response_preview', 'payload',
                      'risk_reasoning', 'attack_narrative', 'confidence',
                      'false_positive_risk', 'proof_status'):
            if not entry.get(field) and f.get(field):
                entry[field] = f.get(field)

    consolidated = []
    for entry in grouped.values():
        urls = _unique_ordered(entry.pop('_affected_urls', []))
        params = _unique_ordered(entry.pop('_parameters', []))
        evidence_items = _unique_ordered(entry.pop('_evidence_items', []))
        entry['affected_urls'] = urls
        entry['parameter'] = ', '.join(params[:6]) if params else entry.get('parameter', '')
        if evidence_items:
            summary = evidence_items[0]
            if len(evidence_items) > 1:
                summary += f'\n\nAdditional evidence examples: ' + ' | '.join(evidence_items[1:4])
            entry['evidence'] = summary
        if urls:
            entry['url'] = urls[0]
        consolidated.append(entry)

    return sorted(
        consolidated,
        key=lambda f: (
            SEVERITY_ORDER.get((f.get('severity') or 'info').lower(), 99),
            f.get('display_name') or f.get('vuln_type') or '',
        ),
    )


# ---------------------------------------------------------------------------
# XML Helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table, color='BFBFBF', sz='4'):
    """Set thin borders on all edges of a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')
        borders.append(el)
    # Remove existing borders if any
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)


def _add_page_number(paragraph):
    """Insert auto-incrementing page number field."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld_char_end)


def _add_hyperlink(paragraph, text, url):
    """Add a clickable URL to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _sanitize(text, max_len=1000):
    """Clean text for DOCX insertion: strip control chars, truncate."""
    if not text:
        return ''
    if isinstance(text, list):
        text = '; '.join(str(t) for t in text)
    text = str(text)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    if len(text) > max_len:
        text = text[:max_len] + ' [...truncated]'
    return text


def _style_exists(doc, style_name):
    try:
        doc.styles[style_name]
        return True
    except Exception:
        return False


def _style_or_none(doc, style_name):
    return style_name if _style_exists(doc, style_name) else None


def _add_bullet(doc, text):
    style = _style_or_none(doc, 'List Bullet')
    if style:
        doc.add_paragraph(text, style=style)
    else:
        doc.add_paragraph(f'- {text}')


def _add_numbered(doc, number, text):
    style = _style_or_none(doc, 'List Number')
    if style:
        doc.add_paragraph(text, style=style)
    else:
        doc.add_paragraph(f'{number}. {text}')


def _clear_document_body(doc):
    """Remove template sample content while preserving section/header/style data."""
    body = doc._element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)


def _new_report_document():
    """Create the report from the default Word template when available."""
    if os.path.exists(DEFAULT_TEMPLATE_PATH):
        doc = Document(DEFAULT_TEMPLATE_PATH)
        _clear_document_body(doc)
        return doc, True
    return Document(), False


def _add_template_cover_page(doc, scan_data):
    """Cover page that uses the uploaded pentest template's named styles."""
    title_style = _style_or_none(doc, 'Subtitle 1')
    meta_style = _style_or_none(doc, 'Doc Number and Version')
    date_style = _style_or_none(doc, 'Version Date')

    title = doc.add_paragraph(style=title_style)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Penetration Testing Assessment Report')
    run.bold = True

    target = scan_data.get('target_url') or 'Target not specified'
    scan_id = scan_data.get('id') or scan_data.get('engagement_id') or 'Generated'
    started = _format_date(scan_data.get('started_at')) or datetime.now().strftime('%B %d, %Y')

    for label, value, style in [
        ('RITM', scan_id, meta_style),
        ('Target', target, meta_style),
        ('Date', started, date_style),
    ]:
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{label}: {value}')
        r.bold = label in {'RITM', 'Target'}

    doc.add_paragraph()
    conf = doc.add_paragraph(style=meta_style)
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('CONFIDENTIAL - FOR AUTHORIZED RECIPIENTS ONLY')
    run.bold = True
    doc.add_page_break()


def _build_curl(finding):
    """Generate a curl command that actually reproduces the attack.

    Precedence for building the curl:
    1. Parse the full HTTP request from `finding.request` if present — this is
       the highest fidelity reproduction because it contains the exact payload,
       headers, and body the agent sent.
    2. Fall back to url + method + param/payload injection.

    Always shell-escape URLs and values so the command is copy-paste safe.
    """
    from urllib.parse import urlparse, urlunparse, parse_qsl, urlencode, quote

    url = finding.get('url', '') or ''
    method = (finding.get('method') or 'GET').upper()
    payload = finding.get('payload', '') or ''
    param = finding.get('parameter', '') or ''
    req = finding.get('request', '') or ''

    if not url and not req:
        return ''

    # ---- 1. Parse the raw request field if present ----
    req_method = None
    req_path = None
    req_host = None
    req_headers = []
    req_body = ''
    if req:
        raw = str(req).replace('\\r\\n', '\n').replace('\\n', '\n').replace('\r\n', '\n')
        # Split into headers and body on first blank line
        if '\n\n' in raw:
            head, req_body = raw.split('\n\n', 1)
        else:
            head, req_body = raw, ''
        lines = [ln for ln in head.split('\n') if ln.strip()]
        for i, line in enumerate(lines):
            if i == 0 and line.split(' ', 1)[0] in ('GET', 'POST', 'PUT', 'DELETE', 'PATCH', 'HEAD', 'OPTIONS'):
                try:
                    req_method, req_path, _ = line.split(' ', 2)
                except ValueError:
                    req_method = line.split(' ', 1)[0]
                continue
            if ':' in line:
                name, _, value = line.partition(':')
                name = name.strip()
                value = value.strip()
                if name.lower() == 'host':
                    req_host = value
                else:
                    req_headers.append((name, value))

    # ---- 2. Build the final URL ----
    final_url = url
    if req_path and req_host:
        # Prefer the request-line path because it contains the actual attack payload
        scheme = urlparse(url).scheme if url else 'https'
        final_url = f"{scheme}://{req_host}{req_path}"
    elif method == 'GET' and param and payload and url and param.lower() not in {'path', 'url', 'route', 'endpoint'}:
        # No request field — inject payload into URL for GET
        parsed = urlparse(url)
        existing = dict(parse_qsl(parsed.query, keep_blank_values=True))
        if existing.get(param) != payload:
            existing[param] = payload
            final_url = urlunparse(parsed._replace(query=urlencode(existing, safe="':\"<>")))

    # ---- 3. Assemble the curl command ----
    def _shell_escape(s):
        # Single-quote wrap and escape embedded single quotes
        return "'" + str(s).replace("'", "'\\''") + "'"

    parts = ['curl -sk']
    final_method = req_method or method
    if final_method and final_method != 'GET':
        parts.append(f'-X {final_method}')

    # Always use a realistic User-Agent so the request resembles the attack
    has_ua = any(h.lower() == 'user-agent' for h, _ in req_headers)
    if not has_ua:
        parts.append('-H ' + _shell_escape('User-Agent: Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'))
    for name, value in req_headers:
        parts.append('-H ' + _shell_escape(f'{name}: {value}'))

    parts.append(_shell_escape(final_url))

    # Body for non-GET requests
    if final_method in ('POST', 'PUT', 'PATCH'):
        body = req_body.strip() if req_body else ''
        if not body and payload:
            body = f'{param}={payload}' if param else payload
        if body:
            parts.append('--data-raw ' + _shell_escape(body))

    return ' \\\n  '.join(parts)


def _retest_steps(f, curl):
    url = f.get('url') or ''
    vuln_type = f.get('vuln_type') or 'finding'
    steps = [
        f'Open the affected URL in an authenticated browser session if the endpoint requires authentication: {url}',
        'Run the provided curl command from a terminal.',
        'Confirm the response matches the evidence shown in this report.',
        f'After remediation, repeat the same request and confirm the {vuln_type} evidence is no longer present.',
    ]
    if not curl:
        steps[1] = 'Send the raw request shown below through Burp Repeater or curl.'
    return steps


def _get_remediation(vuln_type):
    """Look up remediation text by vuln_type, with fuzzy matching."""
    if not vuln_type:
        return DEFAULT_REMEDIATION
    vt = vuln_type.lower().replace('-', '_').replace(' ', '_')
    # Exact match
    if vt in REMEDIATION:
        return REMEDIATION[vt]
    # Substring match
    for key, text in REMEDIATION.items():
        if key in vt or vt in key:
            return text
    # Keyword match
    keywords = {
        'xss': ['xss', 'cross_site_scripting', 'reflected', 'stored', 'dom_based'],
        'sql_injection': ['sql', 'sqli', 'injection'],
        'command_injection': ['command', 'cmd', 'os_command', 'rce'],
        'ssrf': ['ssrf', 'server_side_request'],
        'path_traversal': ['traversal', 'directory', 'lfi', 'file_inclusion'],
        'csrf': ['csrf', 'cross_site_request'],
        'clickjacking': ['clickjack', 'frame', 'ui_redress'],
        'cors': ['cors', 'cross_origin'],
        'jwt': ['jwt', 'json_web_token'],
        'information_disclosure': ['info', 'disclosure', 'leak', 'verbose', 'error_message'],
        'security_headers': ['header', 'missing_header', 'hsts', 'csp'],
        'open_redirect': ['redirect'],
    }
    for key, kws in keywords.items():
        if any(kw in vt for kw in kws):
            return REMEDIATION[key]
    return DEFAULT_REMEDIATION


# ---------------------------------------------------------------------------
# Style Setup
# ---------------------------------------------------------------------------

def _setup_styles(doc):
    """Configure document styles and page layout."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level, (name, size, color) in enumerate([
        ('Heading 1', 20, '1E3A5F'),
        ('Heading 2', 15, '2D3748'),
        ('Heading 3', 12, '4A5568'),
    ], start=1):
        h = doc.styles[name]
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor.from_string(color)
        h.font.bold = True
        h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        h.paragraph_format.space_after = Pt(6)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)


# ---------------------------------------------------------------------------
# Header / Footer
# ---------------------------------------------------------------------------

def _add_header_footer(doc, scan_data):
    """Add header and footer to all pages except the cover."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ''
    run_left = hp.add_run('CONFIDENTIAL')
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = RGBColor(0x9B, 0x1C, 0x31)
    run_left.bold = True
    hp.add_run('\t\t')
    run_right = hp.add_run('White Hat Labs — DAST Report')
    run_right.font.size = Pt(8)
    run_right.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ''
    run_f = fp.add_run('Page ')
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    _add_page_number(fp)
    run_f2 = fp.add_run(f'  |  Scan {scan_data.get("id", "")}')
    run_f2.font.size = Pt(8)
    run_f2.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def _add_cover_page(doc, scan_data):
    """Professional cover page."""
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('PENETRATION TEST REPORT')
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor.from_string('1A1A2E')
    run.bold = True
    run.font.name = 'Calibri'

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Dynamic Application Security Testing (DAST)')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.name = 'Calibri'

    # Divider
    doc.add_paragraph()
    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div.add_run('━' * 50)
    run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Metadata table (no borders)
    target = scan_data.get('target_url', 'N/A')
    started = _format_date(scan_data.get('started_at'))
    status = (scan_data.get('status') or 'N/A').title()

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    meta = [
        ('Target', target),
        ('Date', started),
        ('Scanner', 'White Hat Labs DAST Scanner v3'),
        ('Status', status),
    ]
    for i, (label, value) in enumerate(meta):
        row = table.rows[i]
        # Label cell
        lc = row.cells[0]
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        lr.font.name = 'Calibri'
        lc.width = Inches(1.5)
        # Value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(11)
        vr.font.name = 'Calibri'
        vc.width = Inches(4.5)

    # Remove table borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        borders.append(el)
    existing = tbl_pr.find(qn('w:tblBorders'))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)

    # Confidential notice
    for _ in range(4):
        doc.add_paragraph()
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf.add_run('CONFIDENTIAL — FOR AUTHORIZED RECIPIENTS ONLY')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x9B, 0x1C, 0x31)
    run.bold = True

    doc.add_page_break()


def _add_executive_summary(doc, findings):
    """Executive summary with severity breakdown."""
    doc.add_heading('Executive Summary', level=1)

    total = len(findings)
    counts = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'info': 0}
    for f in findings:
        sev = (f.get('severity') or 'info').lower()
        if sev in counts:
            counts[sev] += 1
        else:
            counts['info'] += 1

    p = doc.add_paragraph()
    p.add_run(
        'A Dynamic Application Security Testing (DAST) assessment was conducted against the target application. '
        'The assessment used the supplied target URLs, captured HTTP context, AI-assisted test selection, and '
        'active verification to identify security weaknesses. Duplicate agent observations are consolidated in '
        'this report by vulnerability family.'
    )

    doc.add_paragraph()

    # Severity table
    table = doc.add_table(rows=len(counts) + 2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # Header
    for i, text in enumerate(['Severity', 'Count']):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows
    for row_idx, (sev, count) in enumerate(counts.items(), start=1):
        row = table.rows[row_idx]
        # Severity cell
        cell_sev = row.cells[0]
        _set_cell_shading(cell_sev, SEVERITY_BG.get(sev, 'FFFFFF'))
        p = cell_sev.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sev.upper())
        run.bold = True
        run.font.color.rgb = SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0))
        run.font.size = Pt(10)
        # Count cell
        cell_cnt = row.cells[1]
        p = cell_cnt.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(count))
        run.font.size = Pt(10)
        if count > 0:
            run.bold = True

    # Total row
    total_row = table.rows[-1]
    _set_cell_shading(total_row.cells[0], 'E5E7EB')
    _set_cell_shading(total_row.cells[1], 'E5E7EB')
    p = total_row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TOTAL')
    run.bold = True
    run.font.size = Pt(10)
    p = total_row.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(total))
    run.bold = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Risk assessment
    doc.add_heading('Risk Assessment', level=2)
    if counts['critical'] > 0:
        risk_text = (
            f'The assessment identified {counts["critical"]} CRITICAL and {counts["high"]} HIGH severity '
            f'vulnerabilities that require immediate remediation. These findings pose a significant risk '
            f'to the confidentiality, integrity, and availability of the application and its data.'
        )
    elif counts['high'] > 0:
        risk_text = (
            f'The assessment identified {counts["high"]} HIGH severity vulnerabilities that should be '
            f'addressed as a priority. While no critical issues were found, these findings could be '
            f'exploited to compromise application security.'
        )
    elif counts['medium'] > 0:
        risk_text = (
            f'The assessment identified {counts["medium"]} MEDIUM severity findings. While these do not '
            f'represent an immediate critical risk, they should be addressed to strengthen the overall '
            f'security posture of the application.'
        )
    elif total > 0:
        risk_text = (
            f'The assessment identified {total} low-severity or informational findings. The overall '
            f'security posture of the application is reasonable, but minor improvements are recommended.'
        )
    else:
        risk_text = (
            'No vulnerabilities were identified during this assessment. The application appears to have '
            'a strong security posture based on the tests performed.'
        )
    doc.add_paragraph(risk_text)


def _add_scope_section(doc, scan_data):
    """Scope and scan details."""
    doc.add_heading('Scope', level=1)

    target = scan_data.get('target_url', 'N/A')
    doc.add_paragraph(f'The following target was assessed:')

    # Target table
    table = doc.add_table(rows=1, cols=2)
    _set_table_borders(table)
    _set_cell_shading(table.rows[0].cells[0], TABLE_HEADER_BG)
    p = table.rows[0].cells[0].paragraphs[0]
    run = p.add_run('Target URL')
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    p = table.rows[0].cells[1].paragraphs[0]
    run = p.add_run(target)
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Scan duration
    started = scan_data.get('started_at')
    ended = scan_data.get('ended_at')
    if started and ended:
        try:
            s = datetime.fromisoformat(started.replace('Z', '+00:00'))
            e = datetime.fromisoformat(ended.replace('Z', '+00:00'))
            dur = e - s
            minutes = int(dur.total_seconds() // 60)
            seconds = int(dur.total_seconds() % 60)
            doc.add_paragraph(f'Scan Duration: {minutes} minutes {seconds} seconds')
        except (ValueError, TypeError):
            pass


def _add_methodology(doc):
    """Methodology section describing the scanning approach."""
    doc.add_heading('Methodology', level=1)

    doc.add_paragraph(
        'The assessment followed a structured five-phase approach designed to maximize coverage '
        'while minimizing false positives:'
    )

    phases = [
        ('Phase 1: Scope Intake',
         'The scanner tested only the supplied target URLs and any Burp/Proxy request context explicitly sent '
         'to Scanner. Broad crawling was not performed from the Scanner tab.'),
        ('Phase 2: Baseline and Classification',
         'Each endpoint was baselined first, then classified by request/response shape to select relevant test '
         'modules such as JWT, API authorization, injection, XXE, SSRF, file handling, and headers.'),
        ('Phase 3: Passive Analysis',
         'Inspection of HTTP responses for information leakage, security misconfigurations, '
         'missing security headers, verbose error messages, and other passive indicators of '
         'security weaknesses without sending any attack payloads.'),
        ('Phase 4: Active Testing',
         'Targeted payload delivery to test for injection vulnerabilities including Cross-Site '
         'Scripting (XSS), SQL Injection, Command Injection, Server-Side Request Forgery (SSRF), '
         'Path Traversal, Clickjacking, JWT weaknesses, and other OWASP Top 10 vulnerabilities.'),
        ('Phase 5: AI-Assisted Verification',
         'Intelligent analysis of scan results using AI models to validate findings, reduce '
         'false positives, and provide context-aware severity assessments based on the specific '
         'application behavior observed during testing.'),
    ]
    for title, desc in phases:
        doc.add_heading(title, level=3)
        doc.add_paragraph(desc)


def _add_code_block(doc, text):
    """Add a gray-background monospace code block using a single-cell table."""
    if not text:
        doc.add_paragraph('N/A')
        return
    text = str(text).replace('\\r\\n', '\n').replace('\\n', '\n')
    text = _sanitize(text, 1500)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, CODE_BLOCK_BG)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # Light border
    _set_table_borders(table, color='CBD5E1', sz='2')


def _images_for_evidence_key(evidence_images, key):
    matches = []
    pattern = re.compile(rf'^{re.escape(key)}(?:$|[\\s._-])', re.I)
    for image in evidence_images or []:
        name = image.get('name') or os.path.basename(image.get('path', ''))
        stem = os.path.splitext(name)[0]
        if pattern.search(stem):
            matches.append(image)
    return sorted(matches, key=lambda x: x.get('name') or '')


def _add_evidence_screenshots(doc, screenshots, evidence_key):
    if not screenshots:
        return
    for idx, image in enumerate(screenshots, 1):
        path = image.get('path')
        if not path or not os.path.exists(path):
            continue
        caption = doc.add_paragraph()
        run = caption.add_run(f'Evidence screenshot {evidence_key}.{idx}: {image.get("name", os.path.basename(path))}')
        run.bold = True
        run.font.size = Pt(9)
        try:
            doc.add_picture(path, width=Inches(6.2))
        except Exception as exc:
            doc.add_paragraph(f'Could not insert screenshot {image.get("name", path)}: {exc}')


# Source-verification status → (human label, hex color)
_STATUS_DISPLAY = {
    'fixed':       ('Remediated in source', '059669'),   # green
    'vulnerable':  ('Vulnerable — not remediated', 'DC2626'),  # red
    'partial':     ('Partially mitigated', 'D97706'),    # amber
    'not_found':   ('Source location not found', '6B7280'),  # gray
    'not_checked': ('Not verified against source', '6B7280'),
}


def _add_source_analysis(doc, f):
    """Render the source-code verification block for a finding.

    No-op in blackbox mode — only renders when source_verifier annotated the
    finding (i.e. the scan was run with --repo-path).
    """
    status = (f.get('remediation_status') or '').lower()
    location = (f.get('source_location') or '').strip()
    evidence = (f.get('remediation_evidence') or '').strip()
    sink = (f.get('source_sink') or '').strip()
    conflict = (f.get('source_conflict') or '').strip()

    # Nothing to show unless source verification actually ran on this finding.
    if not any([status, location, evidence, sink, conflict]):
        return

    doc.add_heading('Source Code Analysis', level=3)

    label, color = _STATUS_DISPLAY.get(status, (status.upper() or 'Unknown', '6B7280'))
    p = doc.add_paragraph()
    run = p.add_run('Status: ')
    run.bold = True
    sr = p.add_run(label)
    sr.bold = True
    sr.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))

    if location:
        p = doc.add_paragraph()
        run = p.add_run('Location: ')
        run.bold = True
        run.font.size = Pt(9.5)
        lr = p.add_run(_sanitize(location, 200))
        lr.font.name = 'Courier New'
        lr.font.size = Pt(9.5)

    if evidence:
        p = doc.add_paragraph()
        run = p.add_run('Finding: ')
        run.bold = True
        p.add_run(_sanitize(evidence, 800))

    # Runtime vs source contradiction — the important one to surface loudly.
    if conflict:
        p = doc.add_paragraph()
        run = p.add_run('⚠ Runtime/source conflict: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        p.add_run(_sanitize(conflict, 600))

    if sink:
        _add_code_block(doc, _sanitize(sink, 600))


def _add_poc(doc, f):
    """Render the generated proof-of-concept for a finding, if present."""
    poc = f.get('poc')
    if isinstance(poc, str):
        try:
            import json as _json
            poc = _json.loads(poc)
        except Exception:
            poc = None
    if not isinstance(poc, dict) or not poc:
        return
    doc.add_heading('Proof of Concept', level=3)
    repro = poc.get('reproduced')
    p = doc.add_paragraph()
    run = p.add_run('Reproduced: ')
    run.bold = True
    rr = p.add_run('Yes' if repro else 'No')
    rr.bold = True
    rr.font.color.rgb = RGBColor(0x05, 0x96, 0x69) if repro else RGBColor(0x6B, 0x72, 0x80)
    if poc.get('impact'):
        p = doc.add_paragraph()
        p.add_run('Impact: ').bold = True
        p.add_run(_sanitize(poc['impact'], 600))
    steps = poc.get('poc_steps') or []
    if steps:
        doc.add_paragraph('Steps:').runs[0].bold = True
        for idx, s in enumerate(steps, 1):
            _add_numbered(doc, idx, _sanitize(str(s), 300))
    if poc.get('poc_request'):
        doc.add_paragraph('Request:').runs[0].bold = True
        _add_code_block(doc, _sanitize(poc['poc_request'], 1200))
    if poc.get('poc_response'):
        doc.add_paragraph('Response (proof):').runs[0].bold = True
        _add_code_block(doc, _sanitize(str(poc['poc_response']), 1200))
    if poc.get('proof'):
        p = doc.add_paragraph()
        p.add_run('Proof: ').bold = True
        p.add_run(_sanitize(poc['proof'], 500))


def _add_detailed_findings(doc, findings, evidence_images=None):
    """Detailed section for each finding with evidence, payload, remediation."""
    doc.add_page_break()
    doc.add_heading('Detailed Findings', level=1)

    if not findings:
        doc.add_paragraph('No vulnerabilities were identified during this assessment.')
        return

    sorted_findings = sorted(findings, key=lambda f: SEVERITY_ORDER.get((f.get('severity') or 'info').lower(), 99))
    severity_counts = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'info': 0}

    for idx, f in enumerate(sorted_findings):
        sev = (f.get('severity') or 'info').lower()
        if sev not in severity_counts:
            sev = 'info'
        severity_counts[sev] += 1
        evidence_key = f.get('evidence_key') or f'{SEVERITY_EVIDENCE_PREFIX.get(sev, "I")}{severity_counts[sev]}'
        screenshots = _images_for_evidence_key(evidence_images or [], evidence_key)
        vuln_type = f.get('vuln_type', 'Unknown')
        display_name = f.get('display_name') or _display_finding_name(vuln_type)

        doc.add_heading(f'{display_name} ({evidence_key})', level=2)

        proof_status = f.get('proof_status')
        if not proof_status:
            confidence = (f.get('confidence') or '').lower()
            proof_status = 'confirmed' if confidence == 'confirmed' else 'unverified_hypothesis'

        meta = doc.add_paragraph()
        for label, value in [
            ('Severity', sev.upper()),
            ('Proof', _sanitize(proof_status, 80).replace('_', ' ').title()),
            ('Confidence', _sanitize(f.get('confidence', 'N/A'), 80).upper()),
            ('Method', f.get('method', 'N/A')),
            ('Parameter', f.get('parameter', 'N/A')),
        ]:
            run = meta.add_run(f'{label}: ')
            run.bold = True
            if label == 'Severity':
                run.font.color.rgb = SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0))
            value_run = meta.add_run(f'{value}   ')
            value_run.font.size = Pt(9.5)

        affected = f.get('affected_urls') or [f.get('url')]
        affected = _unique_ordered(affected)
        if affected:
            doc.add_paragraph('Affected URL(s):').runs[0].bold = True
            for url in affected[:12]:
                p = doc.add_paragraph(style=_style_or_none(doc, 'List Bullet'))
                _add_hyperlink(p, _sanitize(url, 250), url)
            if len(affected) > 12:
                doc.add_paragraph(f'Additional affected URLs omitted from this section: {len(affected) - 12}')

        risk_reasoning = (f.get('risk_reasoning') or '').strip()
        attack_narrative = (f.get('attack_narrative') or '').strip()
        if risk_reasoning:
            p = doc.add_paragraph()
            p.add_run('Risk reasoning: ').bold = True
            p.add_run(_sanitize(risk_reasoning, 900))
        if attack_narrative:
            p = doc.add_paragraph()
            p.add_run('Attack narrative: ').bold = True
            p.add_run(_sanitize(attack_narrative, 900))

        # Evidence
        evidence = f.get('evidence', [])
        if evidence or screenshots:
            doc.add_heading('Evidence', level=3)
            if isinstance(evidence, list):
                for e in evidence:
                    _add_bullet(doc, _sanitize(str(e), 500))
            elif evidence:
                doc.add_paragraph(_sanitize(str(evidence), 500))
            _add_evidence_screenshots(doc, screenshots, evidence_key)

        # Payload
        payload = f.get('payload')
        if payload:
            doc.add_heading('Payload', level=3)
            _add_code_block(doc, payload)

        req = f.get('request')
        curl = f.get('curl_command') or ''
        if not curl:
            curl = _build_curl(f)

        doc.add_heading('Retest Steps', level=3)
        for n, step in enumerate(_retest_steps(f, curl), 1):
            doc.add_paragraph(f'{n}. {step}')

        if curl:
            doc.add_heading('Copy/Paste Curl', level=3)
            _add_code_block(doc, curl)

        if req:
            doc.add_heading('Raw Request', level=3)
            _add_code_block(doc, req)

        # Response
        resp = f.get('response') or f.get('response_preview', '')
        if resp:
            doc.add_heading('Response Evidence', level=3)
            _add_code_block(doc, str(resp)[:2000])

        # Source Code Analysis (only when the scan was run with --repo-path)
        _add_source_analysis(doc, f)

        # Proof of Concept (only when a PoC was generated for this finding)
        _add_poc(doc, f)

        # Remediation
        doc.add_heading('Remediation', level=3)
        # Prefer the source-specific fix (tied to a real file/line) as the
        # headline; fall back to the generic template in pure blackbox mode.
        src_fix = (f.get('source_remediation') or '').strip()
        if src_fix:
            p = doc.add_paragraph()
            run = p.add_run('Source-specific fix: ')
            run.bold = True
            p.add_run(_sanitize(src_fix, 1500))
            gp = doc.add_paragraph()
            gr = gp.add_run('General guidance: ')
            gr.bold = True
            gr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            gp.add_run(_get_remediation(vuln_type))
        else:
            doc.add_paragraph(_get_remediation(vuln_type))

        # Page break between findings (not after the last one)
        if idx < len(sorted_findings) - 1:
            doc.add_paragraph()


def _add_appendix(doc, proxy_entries):
    """Appendix with HTTP traffic log."""
    doc.add_page_break()
    doc.add_heading('HTTP Traffic Log', level=1)

    if not proxy_entries:
        doc.add_paragraph('No HTTP traffic was captured during this assessment.')
        return

    # Summary
    total = len(proxy_entries)
    method_counts = {}
    status_counts = {'2xx': 0, '3xx': 0, '4xx': 0, '5xx': 0, 'other': 0}
    for entry in proxy_entries:
        m = entry.get('method', 'GET')
        method_counts[m] = method_counts.get(m, 0) + 1
        s = entry.get('status', 0)
        if 200 <= s < 300:
            status_counts['2xx'] += 1
        elif 300 <= s < 400:
            status_counts['3xx'] += 1
        elif 400 <= s < 500:
            status_counts['4xx'] += 1
        elif s >= 500:
            status_counts['5xx'] += 1
        else:
            status_counts['other'] += 1

    method_str = ', '.join(f'{m}: {c}' for m, c in sorted(method_counts.items()))
    doc.add_paragraph(f'Total requests captured: {total}')
    doc.add_paragraph(f'Methods: {method_str}')
    status_str = ', '.join(f'{k}: {v}' for k, v in status_counts.items() if v > 0)
    doc.add_paragraph(f'Response codes: {status_str}')

    doc.add_paragraph()

    # Traffic table (capped at 100)
    display = proxy_entries[:100]
    truncated = total > 100

    cols = ['#', 'Method', 'URL', 'Status', 'Length']
    table = doc.add_table(rows=1 + len(display), cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table)

    # Header
    for i, text in enumerate(cols):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8.5)

    # Status code colors
    STATUS_COLORS = {
        2: RGBColor(0x16, 0xA3, 0x4A),  # green
        3: RGBColor(0x25, 0x63, 0xEB),  # blue
        4: RGBColor(0xEA, 0x58, 0x0C),  # orange
        5: RGBColor(0xDC, 0x26, 0x26),  # red
    }

    for idx, entry in enumerate(display):
        row = table.rows[idx + 1]
        status = entry.get('status', 0)
        url = _sanitize(entry.get('url', ''), 80)
        values = [
            str(idx + 1),
            entry.get('method', 'GET'),
            url,
            str(status),
            str(entry.get('length', 0)),
        ]
        for i, val in enumerate(values):
            p = row.cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8)
            # Color status codes
            if i == 3:
                color = STATUS_COLORS.get(status // 100)
                if color:
                    run.font.color.rgb = color
                    run.bold = True

        if idx % 2 == 1:
            for i in range(len(cols)):
                _set_cell_shading(row.cells[i], TABLE_ALT_ROW)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.35)
        row.cells[1].width = Inches(0.7)
        row.cells[2].width = Inches(4.2)
        row.cells[3].width = Inches(0.6)
        row.cells[4].width = Inches(0.7)

    if truncated:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f'Note: Showing first 100 of {total} total requests.')
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def _add_coverage_ledger(doc, coverage_rows, coverage_summary=None):
    """Add dynamic DAST coverage notes and gaps."""
    coverage_rows = coverage_rows or []
    if not coverage_rows:
        return

    doc.add_page_break()
    doc.add_heading('Coverage Ledger', level=1)

    summary = coverage_summary or {}
    counts = summary.get('counts') or {}
    total = int(summary.get('total') or len(coverage_rows))
    urls = int(summary.get('urls') or len({r.get('url') for r in coverage_rows if r.get('url')}))
    tested = int(counts.get('tested') or 0)
    confirmed = int(counts.get('confirmed') or 0)
    partial = int(counts.get('partial') or 0)
    skipped = int(counts.get('skipped') or 0)
    not_applicable = int(counts.get('not_applicable') or counts.get('na') or 0)
    doc.add_paragraph(
        f'Dynamic coverage rows recorded: {total}. URLs with coverage: {urls}. '
        f'Tested: {tested}, confirmed: {confirmed}, partial: {partial}, '
        f'skipped: {skipped}, not applicable: {not_applicable}.'
    )

    gaps = [
        r for r in coverage_rows
        if (r.get('status') or '').lower() in {'partial', 'skipped', 'not_applicable', 'na'}
    ]
    if not gaps:
        doc.add_paragraph('No partial, skipped, or not-applicable coverage rows were recorded.')
        return

    doc.add_heading('Coverage Notes and Gaps', level=2)
    display = gaps[:100]
    cols = ['Status', 'URL', 'Surface', 'Category / Subcategory', 'Reason']
    table = doc.add_table(rows=1 + len(display), cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table)
    for i, text in enumerate(cols):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, TABLE_HEADER_BG)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8.5)

    for idx, row_data in enumerate(display, 1):
        row = table.rows[idx]
        category = row_data.get('category') or row_data.get('area') or ''
        subcategory = row_data.get('subcategory') or ''
        cat_text = f'{category} / {subcategory}' if category and subcategory else (category or subcategory)
        values = [
            (row_data.get('status') or '').replace('_', ' '),
            f"{row_data.get('method') or 'GET'} {row_data.get('url') or ''}",
            row_data.get('surface') or '',
            cat_text,
            row_data.get('reason') or '',
        ]
        for i, val in enumerate(values):
            run = row.cells[i].paragraphs[0].add_run(_sanitize(str(val), 220))
            run.font.size = Pt(8)
        if idx % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, TABLE_ALT_ROW)

    if len(gaps) > len(display):
        doc.add_paragraph(f'Note: Showing first {len(display)} of {len(gaps)} coverage notes/gaps.')


def _extract_uploaded_report(report):
    """Extract readable source material from an uploaded DOCX report."""
    path = report.get('path') if isinstance(report, dict) else str(report)
    name = report.get('name') if isinstance(report, dict) else os.path.basename(path)
    extracted = {'name': name or os.path.basename(path), 'headings': [], 'paragraphs': [], 'tables': []}
    try:
        src = Document(path)
    except Exception as exc:
        extracted['error'] = str(exc)
        return extracted

    seen = set()
    for p in src.paragraphs:
        text = ' '.join((p.text or '').split())
        if not text or text in seen:
            continue
        seen.add(text)
        style = (p.style.name if p.style is not None else '').lower()
        if style.startswith('heading') or any(k in text.lower() for k in ('finding', 'vulnerability', 'executive summary', 'risk')):
            extracted['headings'].append(_sanitize(text, 220))
        else:
            extracted['paragraphs'].append(_sanitize(text, 700))
        if len(extracted['headings']) >= 40 and len(extracted['paragraphs']) >= 80:
            break

    for table in src.tables[:8]:
        rows = []
        for row in table.rows[:30]:
            cells = [' '.join((cell.text or '').split()) for cell in row.cells]
            if any(cells):
                rows.append([_sanitize(c, 220) for c in cells])
        if rows:
            extracted['tables'].append(rows)
    return extracted


def _add_uploaded_source_reports(doc, source_reports):
    """Append extracted content from uploaded report files into the final report."""
    source_reports = source_reports or []
    if not source_reports:
        return

    doc.add_page_break()
    doc.add_heading('Uploaded Source Reports', level=1)
    doc.add_paragraph(
        f'{len(source_reports)} uploaded report file(s) were used as source material for this final report. '
        'The content below is extracted from the uploaded reports and normalized into this template.'
    )

    for idx, report in enumerate(source_reports, 1):
        extracted = _extract_uploaded_report(report)
        doc.add_heading(f'{idx}. {extracted.get("name", "Uploaded report")}', level=2)
        if extracted.get('error'):
            doc.add_paragraph(f'Could not read this report: {extracted["error"]}')
            continue

        if extracted.get('headings'):
            doc.add_heading('Detected Sections and Findings', level=3)
            for heading in extracted['headings'][:30]:
                _add_bullet(doc, heading)

        if extracted.get('paragraphs'):
            doc.add_heading('Key Extracted Narrative', level=3)
            for para in extracted['paragraphs'][:25]:
                doc.add_paragraph(para)

        if extracted.get('tables'):
            doc.add_heading('Extracted Tables', level=3)
            for table_rows in extracted['tables'][:3]:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols)
                _set_table_borders(table, color='E5E7EB', sz='2')
                for r_idx, row_data in enumerate(table_rows):
                    for c_idx in range(cols):
                        text = row_data[c_idx] if c_idx < len(row_data) else ''
                        run = table.rows[r_idx].cells[c_idx].paragraphs[0].add_run(text)
                        run.font.size = Pt(8)
                    if r_idx == 0:
                        for cell in table.rows[r_idx].cells:
                            _set_cell_shading(cell, TABLE_HEADER_BG)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                doc.add_paragraph()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _format_date(iso_str):
    """Format ISO date string to readable format."""
    if not iso_str:
        return 'N/A'
    try:
        dt = datetime.fromisoformat(iso_str.replace('Z', '+00:00'))
        return dt.strftime('%B %d, %Y at %I:%M %p')
    except (ValueError, TypeError):
        return iso_str


# ---------------------------------------------------------------------------
# Main Entry Point
# ---------------------------------------------------------------------------

def generate_docx_report(scan_data, findings, proxy_entries, source_reports=None, evidence_images=None):
    """Generate a professional DOCX penetration test report.

    Args:
        scan_data: dict with target_url, started_at, ended_at, status, crawl_results, id
        findings: list of finding dicts
        proxy_entries: list of proxy entry dicts
        source_reports: uploaded DOCX reports to extract and consolidate
        evidence_images: screenshot files named by finding key, e.g. H1.png, H2.jpg

    Returns:
        io.BytesIO containing the DOCX file
    """
    findings = _consolidate_findings(findings)
    doc, using_template = _new_report_document()
    if using_template:
        _add_template_cover_page(doc, scan_data)
    else:
        _setup_styles(doc)
        _add_header_footer(doc, scan_data)
        _add_cover_page(doc, scan_data)
    _add_executive_summary(doc, findings)
    _add_scope_section(doc, scan_data)
    _add_methodology(doc)
    # Detailed findings only (no separate overview table — it duplicated Section 4)
    _add_detailed_findings(doc, findings, evidence_images=evidence_images or [])
    _add_coverage_ledger(
        doc,
        scan_data.get('checklist_coverage') or [],
        scan_data.get('checklist_coverage_summary') or {},
    )
    _add_uploaded_source_reports(doc, source_reports or [])
    _add_appendix(doc, proxy_entries)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
